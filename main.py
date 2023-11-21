import json
from openai import OpenAI
from pandas import read_csv
from docx import Document
from docx.shared import Pt
import os
from tenacity import retry, wait_chain, wait_fixed

class TextCreator:
    def __init__(   
                    self, 
                    api_key: str,
                    keyword: str,
                    country: str,
                    result_folder_path: str,
                    gpt_version: str,
                    sync_data: bool
                ):
        self.api_key = api_key
        self.keyword = keyword
        self.country = country
        self.result_folder_path = result_folder_path
        self.gpt_version = gpt_version
        self.sync_data = sync_data


    def create_full_text(self):
        self.prestart()
        document = Document()
        document, review_data = self.get_review(document)
        document, info_data = self.get_information(document)
        document, rtp_data = self.get_rtp(document)
        document, jackpot_data = self.get_jackpot(document)
        document, mobile_play_data = self.get_mobile_play(document)
        document, pros_and_cons_data  = self.get_pros_and_cons(document)
        document, faq_data  = self.get_faq(document)
        document.save(os.path.join(self.result_folder_path, f"{self.keyword.replace(' ', '_').lower()}_{self.country.replace(' ', '_').lower()}", f'{self.keyword.replace(' ', '_').lower()}_{self.country.replace(' ', '_').lower()}.docx'))
        if self.sync_data:
            self.syncronize_all_data(review_data, info_data, rtp_data,jackpot_data, mobile_play_data, pros_and_cons_data, faq_data)

    def prestart(self):
        os.makedirs(os.path.join(self.result_folder_path, f"{self.keyword.replace(' ', '_').lower()}_{self.country.replace(' ', '_').lower()}"), exist_ok=True)
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_review(self, document: Document):
        document.add_heading(f"{self.keyword} Review in {self.country}", 1)
        query = f"""
            {self.keyword} Slot Review in {self.country}
            Write intro about slot up to 100 words.
            Describe slot theme, provider, paylines, features, jackpot option, etc.
        """
        generated_text = self.get_message_from_chat(query)
        document.add_paragraph(generated_text)
        data = {}
        data['text'] = generated_text
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_information(self, document: Document):
        document.add_heading(f"H2: {self.keyword} Information", 1)
        query = f"""
            Create a key-value pairs. Key and value separated with ";". Each key-values pair starts from new row. Create with next information on the {self.keyword} slot
            Rating - set random rating from 3.5 to 5 stars. 
            Gamble Feature - write what slot features are available at the {self.keyword} slot
            RTP - write direct RTP or RTP range
            No. of Paylines - write a number
            House Edge - write number or house edge range
            No. of Reels - write a number
            Software Providers - write a software provider name
            Multipliers - write Yes or No
            Autospin Feature - write Yes or No
            Slot Type - write slot type
            Launch Date - write a date with month,date, year separated by coma
            Wild Symbol - write Yes or No 
            Min. Bet - write number or range
            Scatter Symbol - write Yes or No 
            Max. Bet - write the number or range
            Jackpot - write Yes or No 
            Devices - write what device types the slot supports. Available options - Desktop, Mobile, Tablet
            Free Spins - write Yes or No 
        """
        generated_text = self.get_message_from_chat(query)
        result = []
        rows = generated_text.split('\n')
        for row in rows:
            criterea, information = row.split(';')
            criterea = criterea.strip()
            information = information.strip()
            result.append([criterea, information])
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Criterea'
        hdr_cells[0].bold = True
        hdr_cells[1].text = 'Information'
        hdr_cells[1].bold = True
        arr = []
        data = {}
        for criterea, information in result:
            row_cells = table.add_row().cells
            row_cells[0].text = str(criterea)
            row_cells[1].text = str(information)
            pair = {}
            pair["criterea"] = str(criterea)
            pair["info"] = str(information)
            arr.append(pair)
        data['info'] = arr
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_rtp(self, document: Document):
        document.add_heading(f"H2: {self.keyword} RTP and Wager Limits", 1)
        query = f"""
            Write one paragraph of text about {self.keyword} RTP and Wager Limits
            Write in simple words, avoid repetitions and stick to the point.
        """
        generated_text = self.get_message_from_chat(query)
        document.add_paragraph(generated_text)
        data = {}
        data['text'] = generated_text
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_symbols_and_combinations(self, document: Document):
        document.add_heading(f"H2: Symbols and Winning Combinations of {self.keyword}  Slot", 1)
        query = f"""
            Write text about Symbols and Winning Combinations of {self.keyword}  Slot
            Add table with symbols and paytable {self.keyword}  Slot
            Write in simple words, avoid repetitions and stick to the point.
        """
        generated_text = self.get_message_from_chat(query)
        print(generated_text)
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_jackpot(self, document: Document):
        document.add_heading(f"H2: Jackpots and Bonus Features of {self.keyword}  Slot", 1)
        query = f"""
            Firstly write text with key "text" about Jackpots and Bonus Features of {self.keyword}. After that write key "jackpots" and
            create a list with descriptions of Jackpots of {self.keyword} in it. After that write key "bonus_features" and create a list of Bonus Features of {self.keyword} in it. At the end add key "text_2" with small part of Jackpots and Bonus Features. 
            Write in simple words, avoid repetitions and stick to the point. Generate only result in json format.
            """
        generated_text = self.get_message_from_chat(query)
        # print(generated_text)
        data: dict = json.loads(generated_text.replace('`', '').replace('json', ''))
        text = data['text']
        jackpots = data['jackpots']
        bonus_features = data['bonus_features']
        text_2 = data['text_2']
        document.add_paragraph(text)
        document.add_heading('Jackpots', 3)
        for key, jackpot in enumerate(jackpots):
            p = document.add_paragraph(f"{key + 1}. {jackpot}")
            p.paragraph_format.left_indent = Pt(20)
        document.add_paragraph('\n')
        document.add_heading('Bonus Features:', 3)
        for key, bonus_feature in enumerate(bonus_features):
            p = document.add_paragraph(f"{key + 1}. {bonus_feature}")
            p.paragraph_format.left_indent = Pt(20)
        document.add_paragraph(text_2)
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_mobile_play(self, document: Document):
        document.add_heading(f"H2: {self.keyword} Mobile Play", 1)
        query = f"""
            Add one paragraph with an overview of {self.keyword} mobile play experience as "overview" key. Include a list without descriptions of {self.keyword} slot supported devices as 'devices'. Add up to 5 devices that are widely popular in {self.country} in the text. Add small text in the end as "disclaimer".
            Write in simple words, avoid repetitions and stick to the point. Create as json data.
            """
        generated_text = self.get_message_from_chat(query)
        data: dict = json.loads(generated_text.replace('`', '').replace('json', ''))
        text = data['overview']
        devices = data['devices']
        disclaimer = data['disclaimer']
        document.add_paragraph(text)
        document.add_heading('Supported Devices', 3)
        for key, device in enumerate(devices):
            p = document.add_paragraph(f"{key + 1}. {device}")
            p.paragraph_format.left_indent = Pt(20)
        document.add_paragraph(disclaimer)
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_pros_and_cons(self, document: Document):
        document.add_heading(f"H2: Conclusion", 1)
        query = f"""
                    Write a conclusion of {self.keyword} slot review as 'text' key. It must be one small paragraph of text.
                    Include small pros as key 'pros' and cons as key 'cons' of {self.keyword} slot. Add 3-4 pros and 1-2 cons. 
                    Write in simple words, avoid repetitions and stick to the point. Create as json data 
                """
        generated_text = self.get_message_from_chat(query)
        data = json.loads(generated_text.replace('`', '').replace('json', ''))
        text = data['text']
        pros = data['pros']
        cons = data['cons']
        document.add_paragraph(text)
        document.add_heading('Pros and Cons', 3)
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Pros'
        hdr_cells[0].bold = True
        hdr_cells[1].text = 'Cons'
        hdr_cells[1].bold = True
        for key, pros_i in enumerate(pros):
            row_cells = table.add_row().cells
            row_cells[0].text = str(pros_i)
        for key, cons_i in enumerate(cons):
            row = table.rows[key+1].cells
            row[1].text = str(cons_i)
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def get_faq(self, document: Document):
        document.add_heading(f"H2: FAQ", 1)
        query = f"""
        Answer the questions below. Each answer must be one small paragraph from 20 to 40 words. Do not repeat the question in the answer.
        Can I play {self.keyword} for free?
        Are there any {self.keyword} free spins bonuses?
        How do I trigger the free spins bonus round? 
        What is the maximum amount I can win? 
        What is the RTP of {self.keyword}? 
        Is {self.keyword} a volatile game? 
        Write in simple words, avoid repetitions and stick to the point. Firstly write question, then ":" and after that answer for this question. Create as numeric list.
        """
        generated_text = self.get_message_from_chat(query)
        document.add_paragraph(generated_text)
        data = {}
        data['text'] = generated_text
        return document, data
    
    @retry(wait=wait_chain(*[wait_fixed(10) for i in range(6)]))
    def syncronize_all_data(
                            self,
                            review_data: dict,
                            info_data: dict,
                            rtp_data: dict,
                            jackpot_data: dict,
                            mobile_play_data: dict,
                            pros_and_cons_data: dict,
                            faq_data: dict
                        ):
        try:
            all_data = {}
            all_data['review_data'] = review_data
            all_data['info_data'] = info_data
            all_data['rtp_data'] = rtp_data
            all_data['jackpot_data'] = jackpot_data
            all_data['mobile_play_data'] = mobile_play_data
            all_data['pros_and_cons_data'] = pros_and_cons_data
            all_data['faq_data'] = faq_data
            all_data_text = json.dumps(all_data)
            query = f"""
                Check the text below. Remove duplicated information but leave content structure. Rewrite data to synchronize all data through the text.
                The text:

                {all_data_text}

                Return exactly the same json structure.
                """
            generated_data = self.get_message_from_chat(query)
            data = json.loads(generated_data.replace('`', '').replace('json', ''))
            document = Document()
            document.add_heading(f"{self.keyword} Review in {self.country}", 1)
            document.add_paragraph(data['review_data']['text'])

            document.add_heading(f"H2: {self.keyword} Information", 1)
            pairs = data['info_data']['info']
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Criterea'
            hdr_cells[0].bold = True
            hdr_cells[1].text = 'Information'
            hdr_cells[1].bold = True
            for pair in pairs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(pair['criterea'])
                row_cells[1].text = str(pair['info'])

            document.add_heading(f"H2: {self.keyword} RTP and Wager Limits", 1)
            document.add_paragraph(data['rtp_data']['text'])

            document.add_heading(f"H2: Jackpots and Bonus Features of {self.keyword}  Slot", 1)
            document.add_paragraph(data['jackpot_data']['text'])
            document.add_heading('Jackpots', 3)
            for key, jackpot in enumerate(data['jackpot_data']['jackpots']):
                p = document.add_paragraph(f"{key + 1}. {jackpot}")
                p.paragraph_format.left_indent = Pt(20)
            document.add_paragraph('\n')
            document.add_heading('Bonus Features:', 3)
            for key, bonus_feature in enumerate(data['jackpot_data']['bonus_features']):
                p = document.add_paragraph(f"{key + 1}. {bonus_feature}")
                p.paragraph_format.left_indent = Pt(20)
            document.add_paragraph(data['jackpot_data']['text_2'])

            document.add_heading(f"H2: {self.keyword} Mobile Play", 1)
            document.add_paragraph(data['mobile_play_data']['overview'])
            document.add_heading('Supported Devices', 3)
            for key, device in enumerate(data['mobile_play_data']['devices']):
                p = document.add_paragraph(f"{key + 1}. {device}")
                p.paragraph_format.left_indent = Pt(20)
            document.add_paragraph(data['mobile_play_data']['disclaimer'])

            document.add_heading(f"H2: Conclusion", 1)
            document.add_paragraph(data['pros_and_cons_data']['text'])
            document.add_heading('Pros and Cons', 3)
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Pros'
            hdr_cells[0].bold = True
            hdr_cells[1].text = 'Cons'
            hdr_cells[1].bold = True
            for key, pros_i in enumerate(data['pros_and_cons_data']['pros']):
                row_cells = table.add_row().cells
                row_cells[0].text = str(pros_i)
            for key, cons_i in enumerate(data['pros_and_cons_data']['cons']):
                row = table.rows[key+1].cells
                row[1].text = str(cons_i)

            document.add_heading(f"H2: FAQ", 1)
            document.add_paragraph(data['faq_data']['text'])
            document.save(os.path.join(self.result_folder_path, f"{self.keyword.replace(' ', '_').lower()}_{self.country.replace(' ', '_').lower()}", f'{self.keyword.replace(' ', '_').lower()}_{self.country.replace(' ', '_').lower()}(1).docx'))
        except Exception as ex:
            print(ex)
            raise
        
    def get_message_from_chat(self, question: str):
        # try:
        client = OpenAI(api_key=self.api_key)
        models = {
            "GPT-4": "gpt-4", 
            "GPT-4 Turbo": "gpt-4-1106-preview", 
            "GPT-3.5": "gpt-3.5-turbo-1106"
        }
        try:
            response = client.chat.completions.create(
                model=models[self.gpt_version],
                messages=[
                    {
                    "role": "user",
                    "content": question
                    }
                ],
                temperature=1,
                max_tokens=4096,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
        except Exception as ex:
            print(ex)
            raise
        # print(response)
        result = response.choices[0].message.content
        # print(result)
        return result
    
    
    
    def test(self):
        # query = """Generate example of two ordered lists using python-docx library."""
        query = """generate code example with library python-docx. How to get all data from document includes tables"""
        result= self.get_message_from_chat(question=query)
        print(result)